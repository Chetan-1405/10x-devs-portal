from __future__ import annotations

import json
import re
from copy import deepcopy
from datetime import date, datetime
from html import escape
from io import BytesIO
from pathlib import Path
from typing import Any
from uuid import uuid4

from docx import Document
from docx.document import Document as DocumentObject
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import (
    ParagraphStyle,
    getSampleStyleSheet,
)
from reportlab.lib.units import mm
from reportlab.platypus import (
    HRFlowable,
    Paragraph as PdfParagraph,
    SimpleDocTemplate,
    Spacer,
    Table as PdfTable,
    TableStyle,
)

from core.constants import (
    INAUGURATION_DATE,
    LEADERSHIP,
    PORTAL_EMAIL,
    PORTAL_FULL_NAME,
    PORTAL_NAME,
)


# ============================================================
# DOCUMENT COLOURS
# ============================================================

RED = colors.HexColor("#EF3340")
NAVY = colors.HexColor("#172033")
TEXT = colors.HexColor("#344054")
MUTED = colors.HexColor("#667085")
LIGHT_GREY = colors.HexColor("#F4F6F9")
BORDER = colors.HexColor("#DCE3EC")
WHITE = colors.white


# ============================================================
# TEMPLATE PLACEHOLDERS
# ============================================================

OFFER_TEMPLATE_PLACEHOLDERS = {
    "{{FULL_NAME}}",
    "{{REGISTRATION_NUMBER}}",
    "{{APPLICATION_REFERENCE}}",
    "{{CANDIDATE_NUMBER}}",
    "{{ACADEMIC_YEAR}}",
    "{{CLUB_NAME}}",
    "{{CLUB_DOMAIN}}",
    "{{ISSUE_DATE}}",
    "{{DOCUMENT_NUMBER}}",
}


CLUB_DOMAIN_NAMES = {
    "ML Club": "Machine Learning",
    "Computer Vision Club": "Computer Vision",
    "Web Development Club": "Web Development",
}


# ============================================================
# GENERAL HELPERS
# ============================================================

def safe_text(
    value: Any,
    fallback: str = "Not available",
) -> str:
    if value is None:
        return fallback

    cleaned_value = str(value).strip()

    return cleaned_value or fallback


def safe_pdf_text(
    value: Any,
    fallback: str = "Not available",
) -> str:
    return escape(
        safe_text(value, fallback)
    )


def generate_document_number(
    prefix: str,
) -> str:
    current_date = datetime.now().strftime(
        "%Y%m%d"
    )

    short_token = uuid4().hex[:8].upper()

    return (
        f"{prefix}-{current_date}-{short_token}"
    )


def clean_filename_component(
    value: Any,
) -> str:
    cleaned_value = safe_text(
        value,
        "Unknown",
    )

    cleaned_value = re.sub(
        r"[^A-Za-z0-9_-]+",
        "_",
        cleaned_value,
    )

    cleaned_value = re.sub(
        r"_+",
        "_",
        cleaned_value,
    )

    return cleaned_value.strip("_")


def get_offer_letter_filename(
    student: dict[str, Any],
) -> str:
    registration_number = clean_filename_component(
        student.get(
            "registration_number"
        )
    )

    full_name = clean_filename_component(
        student.get(
            "full_name"
        )
    )

    return (
        f"Offer_Letter_"
        f"{registration_number}_"
        f"{full_name}.docx"
    )


def get_club_domain(
    club_name: str,
) -> str:
    return CLUB_DOMAIN_NAMES.get(
        club_name,
        club_name.replace(
            " Club",
            "",
        ),
    )


# ============================================================
# DOCX TEMPLATE VALIDATION
# ============================================================

def iter_block_paragraphs(
    parent: DocumentObject | _Cell,
):
    """
    Yield paragraphs from a document or table cell, including
    paragraphs inside nested tables.
    """

    for paragraph in parent.paragraphs:
        yield paragraph

    for table in parent.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from iter_block_paragraphs(
                    cell
                )


def get_all_docx_paragraphs(
    document: DocumentObject,
) -> list[Paragraph]:
    paragraphs: list[Paragraph] = []

    paragraphs.extend(
        list(
            iter_block_paragraphs(
                document
            )
        )
    )

    for section in document.sections:
        paragraphs.extend(
            list(
                iter_block_paragraphs(
                    section.header
                )
            )
        )

        paragraphs.extend(
            list(
                iter_block_paragraphs(
                    section.footer
                )
            )
        )

        paragraphs.extend(
            list(
                iter_block_paragraphs(
                    section.first_page_header
                )
            )
        )

        paragraphs.extend(
            list(
                iter_block_paragraphs(
                    section.first_page_footer
                )
            )
        )

        paragraphs.extend(
            list(
                iter_block_paragraphs(
                    section.even_page_header
                )
            )
        )

        paragraphs.extend(
            list(
                iter_block_paragraphs(
                    section.even_page_footer
                )
            )
        )

    return paragraphs


def extract_template_text(
    template_bytes: bytes,
) -> str:
    if not template_bytes:
        raise ValueError(
            "The offer-letter template is empty."
        )

    document = Document(
        BytesIO(template_bytes)
    )

    text_parts = [
        paragraph.text
        for paragraph in get_all_docx_paragraphs(
            document
        )
    ]

    return "\n".join(
        text_parts
    )


def validate_offer_letter_template(
    template_bytes: bytes,
) -> dict[str, Any]:
    """
    Validate that the template is a readable DOCX and report
    which placeholders are present or missing.
    """

    try:
        template_text = extract_template_text(
            template_bytes
        )
    except Exception as error:
        raise ValueError(
            "The uploaded file is not a valid DOCX template."
        ) from error

    detected_placeholders = sorted(
        placeholder
        for placeholder in OFFER_TEMPLATE_PLACEHOLDERS
        if placeholder in template_text
    )

    required_placeholders = {
        "{{FULL_NAME}}",
        "{{REGISTRATION_NUMBER}}",
        "{{CLUB_NAME}}",
        "{{CLUB_DOMAIN}}",
        "{{ISSUE_DATE}}",
    }

    missing_required = sorted(
        required_placeholders
        - set(detected_placeholders)
    )

    return {
        "valid": not missing_required,
        "detected_placeholders": (
            detected_placeholders
        ),
        "missing_required_placeholders": (
            missing_required
        ),
    }


# ============================================================
# DOCX PLACEHOLDER REPLACEMENT
# ============================================================

def copy_run_format(
    source_run: Any,
    target_run: Any,
) -> None:
    """
    Copy the common visible formatting properties of a DOCX run.
    """

    target_run.bold = source_run.bold
    target_run.italic = source_run.italic
    target_run.underline = source_run.underline

    target_run.style = source_run.style

    if source_run.font.name:
        target_run.font.name = (
            source_run.font.name
        )

    if source_run.font.size:
        target_run.font.size = (
            source_run.font.size
        )

    if source_run.font.color.rgb:
        target_run.font.color.rgb = (
            source_run.font.color.rgb
        )

    if source_run.font.highlight_color:
        target_run.font.highlight_color = (
            source_run.font.highlight_color
        )

    target_run.font.all_caps = (
        source_run.font.all_caps
    )

    target_run.font.small_caps = (
        source_run.font.small_caps
    )

    target_run.font.strike = (
        source_run.font.strike
    )

    target_run.font.subscript = (
        source_run.font.subscript
    )

    target_run.font.superscript = (
        source_run.font.superscript
    )


def clear_paragraph_runs(
    paragraph: Paragraph,
) -> None:
    paragraph_element = paragraph._element

    for run_element in list(
        paragraph_element.r_lst
    ):
        paragraph_element.remove(
            run_element
        )


def replace_placeholders_in_paragraph(
    paragraph: Paragraph,
    replacements: dict[str, str],
) -> bool:
    """
    Replace placeholders even when Word has divided the placeholder
    across multiple runs.

    The paragraph formatting is retained. Replacement text uses the
    formatting of the first existing run in that paragraph.
    """

    original_text = paragraph.text

    if not original_text:
        return False

    updated_text = original_text

    for placeholder, replacement in replacements.items():
        updated_text = updated_text.replace(
            placeholder,
            replacement,
        )

    if updated_text == original_text:
        return False

    source_run = (
        paragraph.runs[0]
        if paragraph.runs
        else None
    )

    clear_paragraph_runs(
        paragraph
    )

    new_run = paragraph.add_run(
        updated_text
    )

    if source_run is not None:
        copy_run_format(
            source_run,
            new_run,
        )

    return True


def replace_placeholders_in_text_boxes(
    document: DocumentObject,
    replacements: dict[str, str],
) -> int:
    """
    Replace placeholders in XML text nodes, including many Word
    text boxes and shapes that python-docx does not expose directly.
    """

    replacement_count = 0

    document_parts = [
        document.part,
    ]

    for section in document.sections:
        document_parts.extend(
            [
                section.header.part,
                section.footer.part,
                section.first_page_header.part,
                section.first_page_footer.part,
                section.even_page_header.part,
                section.even_page_footer.part,
            ]
        )

    processed_parts: set[int] = set()

    for part in document_parts:
        part_identity = id(part)

        if part_identity in processed_parts:
            continue

        processed_parts.add(
            part_identity
        )

        root_element = part.element

        text_nodes = root_element.xpath(
            ".//w:t"
        )

        for text_node in text_nodes:
            original_text = text_node.text or ""
            updated_text = original_text

            for placeholder, replacement in replacements.items():
                updated_text = updated_text.replace(
                    placeholder,
                    replacement,
                )

            if updated_text != original_text:
                text_node.text = updated_text
                replacement_count += 1

    return replacement_count


def replace_docx_placeholders(
    document: DocumentObject,
    replacements: dict[str, str],
) -> int:
    replacement_count = 0

    for paragraph in get_all_docx_paragraphs(
        document
    ):
        if replace_placeholders_in_paragraph(
            paragraph,
            replacements,
        ):
            replacement_count += 1

    replacement_count += (
        replace_placeholders_in_text_boxes(
            document,
            replacements,
        )
    )

    return replacement_count


# ============================================================
# PERSONALIZED DOCX OFFER LETTER
# ============================================================

def generate_offer_letter_docx(
    student: dict[str, Any],
    template_bytes: bytes,
    issue_date: date | None = None,
    document_number: str | None = None,
) -> tuple[bytes, str, str]:
    """
    Generate a personalized DOCX offer letter.

    Returns:
        generated_file_bytes,
        generated_filename,
        document_number
    """

    if not template_bytes:
        raise ValueError(
            "The offer-letter template is unavailable."
        )

    if safe_text(
        student.get(
            "application_status"
        ),
        "",
    ) != "Selected":
        raise ValueError(
            "Offer letters can be generated only for "
            "students whose status is Selected."
        )

    issue_date = issue_date or date.today()

    document_number = (
        document_number
        or generate_document_number(
            "10X-OFFER"
        )
    )

    club_name = safe_text(
        student.get(
            "club"
        )
    )

    replacements = {
        "{{FULL_NAME}}": safe_text(
            student.get(
                "full_name"
            )
        ),
        "{{REGISTRATION_NUMBER}}": safe_text(
            student.get(
                "registration_number"
            )
        ),
        "{{APPLICATION_REFERENCE}}": safe_text(
            student.get(
                "application_reference"
            )
        ),
        "{{CANDIDATE_NUMBER}}": safe_text(
            student.get(
                "candidate_number"
            )
        ),
        "{{ACADEMIC_YEAR}}": safe_text(
            student.get(
                "study_year"
            )
        ),
        "{{CLUB_NAME}}": club_name,
        "{{CLUB_DOMAIN}}": get_club_domain(
            club_name
        ),
        "{{ISSUE_DATE}}": issue_date.strftime(
            "%d %B %Y"
        ),
        "{{DOCUMENT_NUMBER}}": (
            document_number
        ),
    }

    try:
        document = Document(
            BytesIO(template_bytes)
        )
    except Exception as error:
        raise ValueError(
            "The stored offer-letter template is not "
            "a valid DOCX file."
        ) from error

    replacement_count = replace_docx_placeholders(
        document,
        replacements,
    )

    if replacement_count == 0:
        raise ValueError(
            "No supported placeholders were found in the "
            "offer-letter template."
        )

    output_buffer = BytesIO()

    document.save(
        output_buffer
    )

    output_buffer.seek(0)

    generated_bytes = output_buffer.getvalue()

    filename = get_offer_letter_filename(
        student
    )

    return (
        generated_bytes,
        filename,
        document_number,
    )


# ============================================================
# PDF STYLES
# ============================================================

def get_pdf_styles() -> dict[str, ParagraphStyle]:
    base_styles = getSampleStyleSheet()

    return {
        "document_title": ParagraphStyle(
            "DocumentTitle",
            parent=base_styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=23,
            leading=29,
            textColor=NAVY,
            alignment=TA_CENTER,
            spaceAfter=8,
        ),
        "document_number": ParagraphStyle(
            "DocumentNumber",
            parent=base_styles["Normal"],
            fontName="Helvetica",
            fontSize=9,
            leading=13,
            textColor=MUTED,
            alignment=TA_CENTER,
            spaceAfter=18,
        ),
        "heading": ParagraphStyle(
            "Heading",
            parent=base_styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=14,
            leading=19,
            textColor=NAVY,
            spaceBefore=8,
            spaceAfter=8,
        ),
        "body": ParagraphStyle(
            "Body",
            parent=base_styles["BodyText"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=17,
            textColor=TEXT,
            alignment=TA_LEFT,
            spaceAfter=10,
        ),
        "small": ParagraphStyle(
            "Small",
            parent=base_styles["BodyText"],
            fontName="Helvetica",
            fontSize=8.5,
            leading=13,
            textColor=MUTED,
        ),
        "center": ParagraphStyle(
            "Center",
            parent=base_styles["BodyText"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=17,
            textColor=TEXT,
            alignment=TA_CENTER,
        ),
        "certificate_name": ParagraphStyle(
            "CertificateName",
            parent=base_styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=28,
            leading=34,
            textColor=RED,
            alignment=TA_CENTER,
            spaceBefore=10,
            spaceAfter=10,
        ),
    }


def draw_page_header_footer(
    canvas: Any,
    document: Any,
) -> None:
    page_width, page_height = A4

    canvas.saveState()

    canvas.setFillColor(NAVY)
    canvas.rect(
        0,
        page_height - 24 * mm,
        page_width,
        24 * mm,
        fill=1,
        stroke=0,
    )

    canvas.setFillColor(RED)
    canvas.rect(
        0,
        page_height - 25.5 * mm,
        page_width,
        1.5 * mm,
        fill=1,
        stroke=0,
    )

    canvas.setFont(
        "Helvetica-Bold",
        15,
    )

    canvas.setFillColor(WHITE)
    canvas.drawString(
        18 * mm,
        page_height - 15 * mm,
        PORTAL_NAME,
    )

    canvas.setFont(
        "Helvetica",
        8,
    )

    canvas.setFillColor(
        colors.HexColor(
            "#D7DEEA"
        )
    )

    canvas.drawRightString(
        page_width - 18 * mm,
        page_height - 15 * mm,
        PORTAL_FULL_NAME,
    )

    canvas.setStrokeColor(BORDER)
    canvas.line(
        18 * mm,
        16 * mm,
        page_width - 18 * mm,
        16 * mm,
    )

    canvas.setFont(
        "Helvetica",
        8,
    )

    canvas.setFillColor(MUTED)

    canvas.drawString(
        18 * mm,
        10 * mm,
        f"Official communication • {PORTAL_EMAIL}",
    )

    canvas.drawRightString(
        page_width - 18 * mm,
        10 * mm,
        f"Page {document.page}",
    )

    canvas.restoreState()


def build_pdf(
    story: list[Any],
    title: str,
) -> bytes:
    buffer = BytesIO()

    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=34 * mm,
        bottomMargin=23 * mm,
        title=title,
        author=PORTAL_NAME,
        subject=title,
    )

    document.build(
        story,
        onFirstPage=draw_page_header_footer,
        onLaterPages=draw_page_header_footer,
    )

    buffer.seek(0)

    return buffer.getvalue()


def details_table(
    details: list[tuple[str, Any]],
    styles: dict[str, ParagraphStyle],
) -> PdfTable:
    rows = [
        [
            PdfParagraph(
                f"<b>{safe_pdf_text(label)}</b>",
                styles["body"],
            ),
            PdfParagraph(
                safe_pdf_text(value),
                styles["body"],
            ),
        ]
        for label, value in details
    ]

    table = PdfTable(
        rows,
        colWidths=[
            55 * mm,
            105 * mm,
        ],
        hAlign="LEFT",
    )

    table.setStyle(
        TableStyle(
            [
                (
                    "BACKGROUND",
                    (0, 0),
                    (0, -1),
                    LIGHT_GREY,
                ),
                (
                    "GRID",
                    (0, 0),
                    (-1, -1),
                    0.5,
                    BORDER,
                ),
                (
                    "VALIGN",
                    (0, 0),
                    (-1, -1),
                    "TOP",
                ),
                (
                    "LEFTPADDING",
                    (0, 0),
                    (-1, -1),
                    10,
                ),
                (
                    "RIGHTPADDING",
                    (0, 0),
                    (-1, -1),
                    10,
                ),
                (
                    "TOPPADDING",
                    (0, 0),
                    (-1, -1),
                    8,
                ),
                (
                    "BOTTOMPADDING",
                    (0, 0),
                    (-1, -1),
                    8,
                ),
            ]
        )
    )

    return table


def signature_table(
    styles: dict[str, ParagraphStyle],
) -> PdfTable:
    president = next(
        (
            leader
            for leader in LEADERSHIP
            if leader["role"] == "President"
        ),
        {
            "name": "President",
            "position": PORTAL_NAME,
        },
    )

    coordinator = next(
        (
            leader
            for leader in LEADERSHIP
            if leader["role"] == "Coordinator"
        ),
        {
            "name": "Coordinator",
            "position": PORTAL_NAME,
        },
    )

    table = PdfTable(
        [
            [
                PdfParagraph(
                    (
                        "<br/><br/>"
                        f"<b>{safe_pdf_text(president['name'])}</b>"
                        "<br/>"
                        f"{safe_pdf_text(president['position'])}"
                        "<br/>President, 10x Devs"
                    ),
                    styles["center"],
                ),
                PdfParagraph(
                    (
                        "<br/><br/>"
                        f"<b>{safe_pdf_text(coordinator['name'])}</b>"
                        "<br/>"
                        f"{safe_pdf_text(coordinator['position'])}"
                        "<br/>Coordinator, 10x Devs"
                    ),
                    styles["center"],
                ),
            ]
        ],
        colWidths=[
            80 * mm,
            80 * mm,
        ],
    )

    table.setStyle(
        TableStyle(
            [
                (
                    "VALIGN",
                    (0, 0),
                    (-1, -1),
                    "TOP",
                ),
                (
                    "LINEABOVE",
                    (0, 0),
                    (0, 0),
                    0.7,
                    MUTED,
                ),
                (
                    "LINEABOVE",
                    (1, 0),
                    (1, 0),
                    0.7,
                    MUTED,
                ),
            ]
        )
    )

    return table


# ============================================================
# SUBMISSION RECEIPT PDF
# ============================================================

def generate_submission_receipt_pdf(
    student: dict[str, Any],
    submission: dict[str, Any],
    document_number: str | None = None,
) -> tuple[bytes, str]:
    document_number = (
        document_number
        or generate_document_number(
            "10X-RECEIPT"
        )
    )

    styles = get_pdf_styles()

    selected_tasks = submission.get(
        "selected_tasks",
        [],
    )

    if isinstance(
        selected_tasks,
        str,
    ):
        try:
            selected_tasks = json.loads(
                selected_tasks
            )
        except json.JSONDecodeError:
            selected_tasks = []

    if not isinstance(
        selected_tasks,
        list,
    ):
        selected_tasks = []

    task_rows = [
        [
            str(index),
            PdfParagraph(
                safe_pdf_text(task),
                styles["body"],
            ),
        ]
        for index, task in enumerate(
            selected_tasks,
            start=1,
        )
    ]

    if not task_rows:
        task_rows = [
            [
                "—",
                PdfParagraph(
                    "No specific task information available.",
                    styles["body"],
                ),
            ]
        ]

    task_table = PdfTable(
        [
            [
                PdfParagraph(
                    "<b>No.</b>",
                    styles["body"],
                ),
                PdfParagraph(
                    "<b>Specific Task</b>",
                    styles["body"],
                ),
            ],
            *task_rows,
        ],
        colWidths=[
            16 * mm,
            144 * mm,
        ],
    )

    task_table.setStyle(
        TableStyle(
            [
                (
                    "BACKGROUND",
                    (0, 0),
                    (-1, 0),
                    NAVY,
                ),
                (
                    "TEXTCOLOR",
                    (0, 0),
                    (-1, 0),
                    WHITE,
                ),
                (
                    "GRID",
                    (0, 0),
                    (-1, -1),
                    0.5,
                    BORDER,
                ),
                (
                    "VALIGN",
                    (0, 0),
                    (-1, -1),
                    "TOP",
                ),
                (
                    "LEFTPADDING",
                    (0, 0),
                    (-1, -1),
                    9,
                ),
                (
                    "RIGHTPADDING",
                    (0, 0),
                    (-1, -1),
                    9,
                ),
                (
                    "TOPPADDING",
                    (0, 0),
                    (-1, -1),
                    7,
                ),
                (
                    "BOTTOMPADDING",
                    (0, 0),
                    (-1, -1),
                    7,
                ),
            ]
        )
    )

    submitted_at = (
        submission.get(
            "final_submitted_at"
        )
        or submission.get(
            "submitted_at"
        )
        or "Not available"
    )

    story: list[Any] = [
        Spacer(
            1,
            6 * mm,
        ),
        PdfParagraph(
            "Final Submission Receipt",
            styles["document_title"],
        ),
        PdfParagraph(
            f"Receipt No: {document_number}",
            styles["document_number"],
        ),
        HRFlowable(
            width="100%",
            thickness=1,
            color=RED,
            spaceBefore=2,
            spaceAfter=14,
        ),
        PdfParagraph(
            (
                "This document confirms that the following "
                "student submitted final proof through the "
                "10x Devs portal."
            ),
            styles["body"],
        ),
        Spacer(
            1,
            3 * mm,
        ),
        details_table(
            [
                (
                    "Student Name",
                    student.get(
                        "full_name"
                    ),
                ),
                (
                    "Registration Number",
                    student.get(
                        "registration_number"
                    ),
                ),
                (
                    "Application Reference",
                    student.get(
                        "application_reference"
                    ),
                ),
                (
                    "Candidate Number",
                    student.get(
                        "candidate_number"
                    ),
                ),
                (
                    "Academic Year",
                    student.get(
                        "study_year"
                    ),
                ),
                (
                    "Registered Club",
                    student.get(
                        "club"
                    ),
                ),
                (
                    "Mandatory Task",
                    submission.get(
                        "mandatory_task_name"
                    ),
                ),
                (
                    "Submission State",
                    submission.get(
                        "submission_state",
                        "Final",
                    ),
                ),
                (
                    "Submitted At",
                    submitted_at,
                ),
                (
                    "Application Status",
                    student.get(
                        "application_status"
                    ),
                ),
            ],
            styles,
        ),
        Spacer(
            1,
            7 * mm,
        ),
        PdfParagraph(
            "Submitted Specific Tasks",
            styles["heading"],
        ),
        task_table,
        Spacer(
            1,
            9 * mm,
        ),
        PdfParagraph(
            (
                "The submission has been recorded and will be "
                "reviewed according to the recruitment process. "
                "This receipt does not indicate selection."
            ),
            styles["body"],
        ),
    ]

    return (
        build_pdf(
            story,
            "10x Devs Final Submission Receipt",
        ),
        document_number,
    )


# ============================================================
# SELECTION CERTIFICATE PDF
# ============================================================

def generate_selection_certificate_pdf(
    student: dict[str, Any],
    issue_date: date | None = None,
    document_number: str | None = None,
) -> tuple[bytes, str]:
    issue_date = issue_date or date.today()

    document_number = (
        document_number
        or generate_document_number(
            "10X-CERT"
        )
    )

    styles = get_pdf_styles()

    story: list[Any] = [
        Spacer(
            1,
            10 * mm,
        ),
        PdfParagraph(
            "Certificate of Selection",
            styles["document_title"],
        ),
        PdfParagraph(
            document_number,
            styles["document_number"],
        ),
        Spacer(
            1,
            7 * mm,
        ),
        PdfParagraph(
            "This is to certify that",
            styles["center"],
        ),
        PdfParagraph(
            safe_pdf_text(
                student.get(
                    "full_name"
                ),
                "Student",
            ),
            styles["certificate_name"],
        ),
        PdfParagraph(
            "has been selected as a member of",
            styles["center"],
        ),
        Spacer(
            1,
            5 * mm,
        ),
        PdfParagraph(
            (
                f"<b>"
                f"{safe_pdf_text(student.get('club'))}"
                f"</b>"
            ),
            ParagraphStyle(
                "ClubCertificate",
                parent=styles["center"],
                fontName="Helvetica-Bold",
                fontSize=19,
                leading=24,
                textColor=NAVY,
            ),
        ),
        Spacer(
            1,
            8 * mm,
        ),
        PdfParagraph(
            (
                "under the 10x Devs student technical community, "
                "in recognition of the student's submitted work, "
                "technical interest and participation in the "
                "recruitment process."
            ),
            styles["center"],
        ),
        Spacer(
            1,
            10 * mm,
        ),
        details_table(
            [
                (
                    "Registration Number",
                    student.get(
                        "registration_number"
                    ),
                ),
                (
                    "Application Reference",
                    student.get(
                        "application_reference"
                    ),
                ),
                (
                    "Academic Year",
                    student.get(
                        "study_year"
                    ),
                ),
                (
                    "Date of Issue",
                    issue_date.strftime(
                        "%d %B %Y"
                    ),
                ),
            ],
            styles,
        ),
        Spacer(
            1,
            22 * mm,
        ),
        signature_table(
            styles
        ),
        Spacer(
            1,
            7 * mm,
        ),
        PdfParagraph(
            (
                f"10x Devs was officially inaugurated on "
                f"{INAUGURATION_DATE}."
            ),
            styles["small"],
        ),
    ]

    return (
        build_pdf(
            story,
            "10x Devs Certificate of Selection",
        ),
        document_number,
    )